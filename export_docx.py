from docx import Document


def save_as_docx(minutes: dict, filename: str):
    doc = Document()
    doc.add_heading('Meeting Minutes', 0)

    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()  # blank line

    doc.save(filename)

